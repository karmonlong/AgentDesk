"""
文档处理工具函数
用于各种文档操作的工具函数
"""

import os
import re
from typing import Dict, List, Tuple, Optional
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_operation_prompt(operation: str, content: str, instruction: str = "") -> str:
    """
    根据操作类型生成相应的提示词

    Args:
        operation: 操作类型: summarize/generate/convert/extract_table
        content: 文档内容
        instruction: 额外指示

    Returns:
        AI提示词
    """
    # 限制内容长度（避免token超限）
    content = content[:8000]

    prompts = {
        "summarize": f"""
请对以下文档进行总结：

文档内容：
{content}

要求：
1. 提取3-5个核心要点
2. 生成简洁准确的摘要（300-500字）
3. 保持客观中立
4. 使用清晰的标题和列表
5. 保留关键数据和事实
{instruction}
""",

        "generate": f"""
请基于以下文档内容生成：

文档内容：
{content}

生成要求：
1. 请根据上下文和用户需求生成合适的文档
2. 可能是：报告、邮件、提案、说明文档等
3. 内容应专业、结构清晰、逻辑严谨
4. 长度适中，符合商务标准
{instruction}
""",

        "convert": f"""
请将以下文档转换为更结构化的格式：

文档内容：
{content}

转换要求：
1. 转换为更容易阅读的格式（如Markdown）
2. 保留所有重要信息和结构
3. 使用标题、列表、表格等结构化元素
4. 如有表格数据，用Markdown表格格式
5. 如有列表，使用有序/无序列表
{instruction}
""",

        "extract_table": f"""
请从以下文档中提取表格和数据：

文档内容：
{content}

提取要求：
1. 识别所有表格和数据
2. 用Markdown表格格式展示
3. 如有多个表格，分别标记
4. 提取表格标题和说明
5. 保留数据的准确性
{instruction}
""",

        "extract_key_points": f"""
请从以下文档中提取关键信息点：

文档内容：
{content}

提取要求：
1. 找出所有重要事实和数据
2. 分类整理（如：时间、地点、人物、数字）
3. 使用清晰的列表格式
4. 注明信息来源位置（如第几段）
5. 提取潜在的行动项或待办事项
{instruction}
""",

        "analyze": f"""
请对以下文档进行深入分析：

文档内容：
{content}

分析要求：
1. 分析文档的主题和目的
2. 评估作者的观点和立场
3. 识别论点和支持证据
4. 指出潜在的逻辑问题或缺失信息
5. 提供你的专业见解
{instruction}
"""
    }

    # 如果没有匹配的操作，使用通用提示词
    if operation not in prompts:
        prompts[operation] = f"""
请对以下文档执行 '{operation}' 操作：

文档内容：
{content}

用户指示：
{instruction}
"""

    return prompts[operation]


def extract_email_addresses(text: str) -> List[str]:
    """
    从文本中提取邮箱地址

    Args:
        text: 输入文本

    Returns:
        邮箱地址列表
    """
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    return re.findall(email_pattern, text)


def extract_phone_numbers(text: str) -> List[str]:
    """
    从文本中提取电话号码

    Args:
        text: 输入文本

    Returns:
        电话号码列表
    """
    # 匹配常见格式的电话号码
    phone_pattern = r'\b(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}\b'
    return re.findall(phone_pattern, text)


def extract_dates(text: str) -> List[str]:
    """
    从文本中提取日期

    Args:
        text: 输入文本

    Returns:
        日期列表
    """
    # 匹配多种日期格式
    date_patterns = [
        r'\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b',  # YYYY-MM-DD, YYYY/MM/DD
        r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b',  # MM/DD/YY, DD/MM/YYYY
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b',  # Month DD, YYYY
        r'\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b'  # DD Month YYYY
    ]

    dates = []
    for pattern in date_patterns:
        dates.extend(re.findall(pattern, text, re.IGNORECASE))

    return list(set(dates))  # 去重


def extract_amounts(text: str) -> List[str]:
    """
    从文本中提取金额

    Args:
        text: 输入文本

    Returns:
        金额列表
    """
    # 匹配货币格式
    amount_pattern = r'[$€£¥]?\s*\d+(?:[,\.]\d{3})*(?:[,\.]\d{2})?\s*(?:USD|EUR|GBP|CNY|美元|欧元|英镑|人民币|元)?'
    return re.findall(amount_pattern, text)


def split_into_sections(text: str, min_length: int = 500) -> List[Tuple[str, str]]:
    """
    将长文本分割成多个部分

    Args:
        text: 输入文本
        min_length: 每个部分的最小长度

    Returns:
        (部分标题, 部分内容) 的元组列表
    """
    sections = []

    # 按章节标题分割（中文和英文）
    # 匹配常见的标题格式
    heading_patterns = [
        r'^.+?:\s*$',  # 中文标题格式
        r'^#+\s+.+$',  # Markdown 标题
        r'^\d+\.\s+.+$',  # 编号标题
        r'^第[一二三四五六七八九十]+章\s+.+$',  # 中文章节
        r'^第\d+章\s+.+$'
    ]

    lines = text.split('\n')
    current_section = []
    current_title = "引言"

    for i, line in enumerate(lines):
        line = line.strip()

        # 检查是否是标题行
        is_heading = False
        for pattern in heading_patterns:
            if re.match(pattern, line, re.MULTILINE):
                is_heading = True
                break

        # 如果是标题，保存当前部分并新建
        if is_heading and current_section:
            section_text = '\n'.join(current_section)
            if len(section_text) >= min_length:
                sections.append((current_title, section_text))
                current_section = []
                current_title = line
            else:
                current_section.append(line)
        else:
            current_section.append(line)

        # 如果已经有很多部分，剩余内容合并到最后
        if len(sections) >= 10 and i > len(lines) * 0.7:
            current_section.extend(lines[i + 1:])
            break

    # 添加最后一个部分
    if current_section:
        section_text = '\n'.join(current_section)
        if len(section_text) >= min_length or not sections:
            sections.append((current_title, section_text))

    return sections


def calculate_statistics(text: str) -> Dict[str, any]:
    """
    计算文本统计信息

    Args:
        text: 输入文本

    Returns:
        统计信息字典
    """
    stats = {
        "total_chars": len(text),
        "total_words": len(text.split()),
        "total_lines": len(text.split('\n')),
        "paragraph_count": len([p for p in text.split('\n\n') if p.strip()]),
        "email_count": len(extract_email_addresses(text)),
        "phone_count": len(extract_phone_numbers(text)),
        "date_count": len(extract_dates(text)),
        "table_count": len(re.findall(r'\|', text)) // 2,  # 估算表格数量
    }

    return stats


def generate_markdown_table(headers: List[str], rows: List[List[str]]) -> str:
    """
    生成 Markdown 格式的表格

    Args:
        headers: 表头列表
        rows: 行数据列表

    Returns:
        Markdown 表格字符串
    """
    if not headers or not rows:
        return ""

    # 创建表头行
    table = [" | ".join(headers)]

    # 创建分隔行
    table.append(" | ".join(["---"] * len(headers)))

    # 添加数据行
    for row in rows:
        # 确保每行有足够的列
        row = row[:len(headers)]  # 限制列数
        row.extend([""] * (len(headers) - len(row)))  # 填充空列
        table.append(" | ".join(str(cell) for cell in row))

    return "\n".join(table)


def create_summary_card(title: str, data: Dict[str, str], emoji: str = "📊") -> str:
    """
    创建摘要卡片（Markdown格式）

    Args:
        title: 卡片标题
        data: 键值对数据
        emoji: 卡片图标

    Returns:
        Markdown 格式的卡片
    """
    lines = [f"### {emoji} {title}", "", "| 项目 | 内容 |", "|------|------|"]

    for key, value in data.items():
        lines.append(f"| {key} | {value} |")

    lines.append("")
    return "\n".join(lines)


def validate_output(output: str, min_length: int = 50, max_length: int = 50000) -> Tuple[bool, str]:
    """
    验证输出质量

    Args:
        output: 生成的内容
        min_length: 最小长度
        max_length: 最大长度

    Returns:
        (是否通过, 原因)
    """
    length = len(output)

    if length < min_length:
        return False, f"内容太短（{length} 字符），可能未正确处理"

    if length > max_length:
        return False, f"内容过长（{length} 字符），可能需要截断"

    if "---CONFIDENCE_LOW---" in output:
        return False, "置信度低，需要人工审核"

    if "I don't have access" in output or "I'm not sure" in output:
        return False, "模型不确定如何处理，需要澄清"

    # 检查是否是重复内容
    words = output.split()
    unique_words_ratio = len(set(words)) / len(words) if words else 0
    if unique_words_ratio < 0.3 and len(words) > 100:
        return False, "检测到重复内容，质量可能不高"

    return True, "通过验证"


__all__ = [
    "get_operation_prompt",
    "extract_email_addresses",
    "extract_phone_numbers",
    "extract_dates",
    "extract_amounts",
    "split_into_sections",
    "calculate_statistics",
    "generate_markdown_table",
    "create_summary_card",
    "validate_output"
]


def markdown_to_docx(md: str, output_path: str) -> str:
    """
    将 Markdown 文本导出为 DOCX（基础支持：标题、段落、列表、代码块）

    Args:
        md: Markdown 文本
        output_path: 输出路径（.docx）

    Returns:
        生成的文件路径
    """
    doc = Document()
    # 标题页简单样式
    lines = (md or "").splitlines()
    in_code = False

    for raw in lines:
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            in_code = not in_code
            if in_code:
                p = doc.add_paragraph()
                p.add_run("代码片段:")
            else:
                p = doc.add_paragraph("")
            continue

        if in_code:
            p = doc.add_paragraph(line)
            p.style = doc.styles['Quote'] if 'Quote' in [s.name for s in doc.styles] else None
            continue

        # 标题
        if line.startswith("###### "):
            doc.add_heading(line[7:], level=6)
            continue
        if line.startswith("##### "):
            doc.add_heading(line[6:], level=5)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            continue
        if line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # 列表
        if re.match(r"^\s*[-\*]\s+", line):
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(re.sub(r"^\s*\d+\.\s+", "", line), style='List Number')
            continue

        # 普通段落
        doc.add_paragraph(line)

    doc.save(output_path)
    return output_path
