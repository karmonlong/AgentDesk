"""
文件处理工具函数
支持多种文档格式的读取和写入
"""

import os
import magic
import PyPDF2
from docx import Document
from docx.document import Document as DocType
import openpyxl
import pandas as pd
from typing import Optional, Tuple, List
import json


def detect_file_type(file_path: str) -> str:
    """
    检测文件类型

    Args:
        file_path: 文件路径

    Returns:
        文件类型字符串: txt/pdf/docx/xlsx/md/json/csv/unknown
    """
    try:
        # 使用 python-magic 检测文件类型
        mime = magic.Magic(mime=True)
        file_mime = mime.from_file(file_path)

        # 根据 MIME 类型映射文件类型
        mime_to_type = {
            "text/plain": "txt",
            "text/x-python": "txt",  # Python文件也当文本处理
            "application/pdf": "pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
            "application/vnd.ms-excel": "xlsx",
            "text/markdown": "md",
            "application/json": "json",
            "text/csv": "csv"
        }

        file_type = mime_to_type.get(file_mime, "unknown")

        # 备用检测：通过文件扩展名
        if file_type == "unknown":
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            ext_to_type = {
                ".txt": "txt",
                ".md": "md",
                ".py": "txt",
                ".pdf": "pdf",
                ".docx": "docx",
                ".xlsx": "xlsx",
                ".csv": "csv",
                ".json": "json"
            }
            file_type = ext_to_type.get(ext, "unknown")

        return file_type

    except Exception as e:
        print(f"检测文件类型失败: {e}")
        return "unknown"


def read_text_file(file_path: str, encoding: str = "utf-8") -> str:
    """
    读取文本文件

    Args:
        file_path: 文件路径
        encoding: 文件编码

    Returns:
        文件内容
    """
    with open(file_path, "r", encoding=encoding, errors="ignore") as f:
        return f.read()


def read_pdf_file(file_path: str) -> str:
    """
    读取 PDF 文件

    Args:
        file_path: PDF文件路径

    Returns:
        提取的文本内容
    """
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text_content = []

            print(f"   正在读取 PDF，共 {len(reader.pages)} 页...")

            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_content.append(f"=== 第 {i + 1} 页 ===\n{text}\n")
                except Exception as e:
                    print(f"   警告: 第 {i + 1} 页提取失败: {e}")

            return "\n".join(text_content)

    except Exception as e:
        print(f"   读取 PDF 失败: {e}")
        return ""


def read_docx_file(file_path: str) -> str:
    """
    读取 Word 文档 (docx)

    Args:
        file_path: Word文档路径

    Returns:
        提取的文本内容
    """
    try:
        doc = Document(file_path)
        text_content = []

        print(f"   正在读取 Word 文档...")

        # 读取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # 读取表格
        if doc.tables:
            text_content.append("\n=== 表格内容 ===")
            for i, table in enumerate(doc.tables):
                text_content.append(f"\n--- 表格 {i + 1} ---")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content.append(" | ".join(row_text))

        return "\n".join(text_content)

    except Exception as e:
        print(f"   读取 Word 文档失败: {e}")
        return ""


def read_excel_file(file_path: str) -> str:
    """
    读取 Excel 文件

    Args:
        file_path: Excel文件路径

    Returns:
        提取的文本内容（表格形式）
    """
    try:
        workbook = openpyxl.load_workbook(file_path, read_only=True)
        text_content = []

        print(f"   正在读取 Excel，共 {len(workbook.sheetnames)} 个工作表...")

        for sheet_name in workbook.sheetnames:
            text_content.append(f"\n=== 工作表: {sheet_name} ===")
            sheet = workbook[sheet_name]

            # 读取前100行
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i >= 100:  # 限制行数
                    text_content.append("... (剩余行数省略)")
                    break

                # 过滤空行
                row_data = [str(cell) if cell is not None else "" for cell in row]
                if any(cell.strip() for cell in row_data):
                    text_content.append(" | ".join(row_data))

        return "\n".join(text_content)

    except Exception as e:
        print(f"   读取 Excel 失败: {e}")
        return ""


def read_csv_file(file_path: str) -> str:
    """
    读取 CSV 文件

    Args:
        file_path: CSV文件路径

    Returns:
        提取的文本内容
    """
    try:
        df = pd.read_csv(file_path, nrows=100)  # 限制行数
        print(f"   正在读取 CSV，共 {len(df)} 行 x {len(df.columns)} 列...")
        return df.to_string(index=False)

    except Exception as e:
        print(f"   读取 CSV 失败: {e}")
        return ""


def read_json_file(file_path: str) -> str:
    """
    读取 JSON 文件

    Args:
        file_path: JSON文件路径

    Returns:
        格式化后的 JSON 字符串
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # 如果是列表，只显示前10项摘要
        if isinstance(data, list):
            preview = {
                "type": "list",
                "length": len(data),
                "preview": data[:10]
            }
            return json.dumps(preview, ensure_ascii=False, indent=2)

        # 如果是字典，直接格式化
        return json.dumps(data, ensure_ascii=False, indent=2)

    except Exception as e:
        print(f"   读取 JSON 失败: {e}")
        # 作为文本文件读取
        return read_text_file(file_path)


def read_file(file_path: str, file_type: Optional[str] = None) -> str:
    """
    通用文件读取函数，自动检测类型并读取

    Args:
        file_path: 文件路径
        file_type: 文件类型（如果已知）

    Returns:
        文件内容字符串
    """
    if file_type is None:
        file_type = detect_file_type(file_path)

    print(f"   使用 {file_type} 格式读取...")

    # 根据文件类型调用相应函数
    readers = {
        "txt": read_text_file,
        "pdf": read_pdf_file,
        "docx": read_docx_file,
        "xlsx": read_excel_file,
        "md": read_text_file,
        "json": read_json_file,
        "csv": read_csv_file
    }

    reader = readers.get(file_type, read_text_file)
    return reader(file_path)


def save_file(file_path: str, content: str) -> bool:
    """
    保存文件

    Args:
        file_path: 保存路径
        content: 文件内容

    Returns:
        是否成功
    """
    try:
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        print(f"✅ 文件保存成功: {file_path}")
        return True

    except Exception as e:
        print(f"❌ 文件保存失败: {e}")
        return False


def get_file_info(file_path: str) -> dict:
    """
    获取文件信息

    Args:
        file_path: 文件路径

    Returns:
        包含文件信息的字典
    """
    if not os.path.exists(file_path):
        return {"error": "文件不存在"}

    stat = os.stat(file_path)
    file_type = detect_file_type(file_path)

    return {
        "path": file_path,
        "name": os.path.basename(file_path),
        "size": stat.st_size,
        "size_mb": round(stat.st_size / (1024 * 1024), 2),
        "type": file_type,
        "last_modified": stat.st_mtime
    }


def list_supported_formats() -> List[str]:
    """
    返回支持的文件格式列表

    Returns:
        格式列表
    """
    return [
        ".txt - 文本文件",
        ".md - Markdown文件",
        ".pdf - PDF文档",
        ".docx - Word文档",
        ".xlsx - Excel表格",
        ".csv - CSV数据文件",
        ".json - JSON数据文件",
        ".py - Python代码（当文本处理）"
    ]


# 导出工具函数
__all__ = [
    "detect_file_type",
    "read_file",
    "read_text_file",
    "read_pdf_file",
    "read_docx_file",
    "read_excel_file",
    "read_csv_file",
    "read_json_file",
    "save_file",
    "get_file_info",
    "list_supported_formats"
]
